# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING

FONT = "맑은 고딕"
NAVY = RGBColor(0x00, 0x33, 0x66)

doc = Document()
st = doc.styles["Normal"]
st.font.name = FONT; st.font.size = Pt(11)
st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

def kfont(run, size=11, bold=False, color=None):
    run.font.name = FONT; run.font.size = Pt(size); run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:eastAsia"), FONT); rf.set(qn("w:ascii"), FONT); rf.set(qn("w:hAnsi"), FONT)
    if color: run.font.color.rgb = color

def title(text, size=18):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(4); pf.space_after = Pt(14); pf.line_spacing = 1.2
    kfont(p.add_run(text), size, True, NAVY)

def heading(text, size=13, before=14):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(8)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.3
    kfont(p.add_run(text), size, True, NAVY)

def body(text, size=11, after=6):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.7
    kfont(p.add_run(text), size)

def label_body(label, text, size=11):
    """label bold + body, parse **bold** inline markers"""
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.7
    if label:
        kfont(p.add_run(label), size, True)
    # split bold segments **..**
    import re
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for seg in parts:
        if seg.startswith("**") and seg.endswith("**"):
            kfont(p.add_run(seg[2:-2]), size, True)
        elif seg:
            kfont(p.add_run(seg), size)

def bullet(text, indent=14):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.left_indent = Pt(indent); pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.6
    kfont(p.add_run("• "), 11)
    import re
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for seg in parts:
        if seg.startswith("**") and seg.endswith("**"):
            kfont(p.add_run(seg[2:-2]), 11, True)
        elif seg:
            kfont(p.add_run(seg), 11)

def kv(k, v):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.left_indent = Pt(14); pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.5
    kfont(p.add_run("• "), 11)
    kfont(p.add_run(f"{k} : "), 11, True)
    kfont(p.add_run(v), 11)

def divider():
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    kfont(p.add_run("─" * 28), 9, color=RGBColor(0xAA, 0xAA, 0xAA))

# ============ 자기소개서 ============
title("현대모비스 Global OE 조향/현가 영업 — 자기소개서 및 경력기술서")

heading("[문항 1] 지원동기와 입사 후 회사에서 이루고 싶은 꿈은 무엇인가요? (972자)", before=8)
body("저는 3년간 폴리우레탄 원료 PPG를 글로벌 B2B 시장에 팔아온 소재 영업맨입니다. 어느 날 제가 판 PPG의 최종 행선지를 추적하다, 그것이 자동차 현가장치의 부싱·마운트·방진고무가 된다는 사실을 알았습니다. 그 순간 제 커리어의 방향이 선명해졌습니다. 소재의 원가 구조와 글로벌 수급을 아는 제가 그 소재로 만든 부품을 완성차 업체에 직접 공급하는 OE 영업을 한다면, 누구보다 설득력 있는 원가 협상을 할 수 있다는 확신이 들었습니다. 3년간 소재의 가장 앞단에서 국제 시황을 읽어온 경험은, 부품이 완성차에 실리는 마지막 단계에서 누구도 쉽게 흉내 내기 어려운 차별점이 될 것입니다.")
body("현대모비스를 택한 이유는 직무의 구조 때문입니다. 글로벌 Top 10 부품사인 현대모비스의 OE 영업은 단순 ‘세일즈’가 아니라, RFQ 대응·원가분석·견적·가격협상부터 수주 후 개발 총괄, 양산 안정화까지 프로젝트 전 사이클을 주도하는 ‘비즈니스 오너’입니다. 이는 제가 PPG에서 시장조사부터 견적·협상·계약·납품·클레임 대응까지 혼자 돌려온 B2B 풀사이클과 구조적으로 같습니다. 게다가 조향·현가는 R-MDPS, 전자식 서스펜션으로 진화 중입니다. 전동화 전환의 한복판에서 영업할 수 있다는 점이 저를 끌어당겼습니다.")
body("입사 후 꿈은 단계적입니다. 단기(1~3년)에는 현대차·기아 해외공장 담당으로 OE 영업 사이클을 체화하고, 영어·일본어를 활용해 아시아·북미 신규 OE 고객 발굴에 도전하겠습니다. 중기(5년)에는 전동화 플랫폼 전담 글로벌 OE 영업 전문가로 성장해 현대모비스의 해외 OE 매출 비중을 확대하는 데 기여하겠습니다. 특히 일본계 완성차와 북미 시장은 제 언어 역량과 글로벌 협상 경험이 즉시 통하는 무대라 확신합니다. 장기적으로는 ‘소재를 아는 부품 영업맨’이라는 차별화된 정체성으로 글로벌 OE 조직을 이끄는 리더가 되겠습니다. 소재에서 부품으로 가치사슬을 거슬러 오르는 제 여정의 종착지가 바로 현대모비스입니다.")
divider()

heading("[문항 2] 지원한 직무를 위해 필요한 역량은 무엇이라고 생각하며, 이 역량을 갖추기 위한 노력이나 자신만의 특별한 경험을 작성해 주세요. (964자)")
label_body("", "**①논리적 분석 능력.** 인도 시장은 정보가 닫혀 있었습니다. 막연한 영업 대신 데이터를 택했습니다. 수출입 통관 데이터 10만 건을 파이썬으로 수집·전처리했고, 경쟁사의 가격대·물량·유통경로·주요 항구를 역추적했습니다. 처음엔 데이터 노이즈로 분석이 어그러졌지만, HS코드 필터 기준을 세 번 고친 끝에 경쟁사가 놓친 틈새 항구를 찾아 진입 루트를 확보했습니다. 같은 분석 역량으로 한 달 걸리던 감사 자료 준비를 AI와 VBA로 자동화해 3일로 단축하기도 했습니다.")
label_body("", "**②진취적 추진력.** 일본에서 경쟁사 철수 정보를 입수했습니다. ‘기회다’로 그치지 않고 그 경쟁사의 고객 리스트를 확보해 품질 테스트 데이터와 납기 경쟁력을 담은 제안서를 보냈습니다. 첫 달 3곳에서 거절당했고, 사유를 뜯어보니 가격이 아닌 ‘공급 안정성’ 우려였습니다. 제안 논리를 안정성 중심으로 피벗했고, 6개월 만에 물량을 42톤에서 84톤으로 두 배 늘렸습니다.")
label_body("", "**③조율 협상력.** 이란-미국 갈등으로 중동발 PO 공급이 흔들리자, 글로벌 PO 플랜트 가동률 데이터로 ‘3개월 내 공급 부족’ 보고서를 만들어 고객사에 선제 공유했습니다. ‘지금 장기 계약이 고객사에 유리하다’는 논리로 톤당 600달러 인상 조건으로 갱신했습니다. 반대로 브라질 ITAPOA의 톤당 30달러 인하 요구는 시황 데이터로 방어했습니다. 올릴 때와 지킬 때를 모두 겪은 협상가입니다.")
label_body("", "**④외국어·글로벌 커뮤니케이션.** 영어(OPIc AL)로 협상하고 4개 문화권 거래처를 동시에 관리합니다. 일본 고객과는 호렌소·장기 신뢰 기반 관행을 이해하는 태도로 신뢰를 쌓았습니다. 이 모든 역량의 토대는 학습력입니다. 화학 비전공자로 입사해 매일 아침 2시간씩 3개월을 독학해 1년 만에 팀 매출 20%를 책임지는 담당자가 됐습니다. 우대사항인 비즈니스 영어와 일본어를 모두 충족한다는 점도 강점입니다.")
divider()

heading("[문항 3] 다른 지원자 대비 본인만의 ‘차별화된 강점’과 ‘보완해야 할 약점’에 대해 사례를 들어 구체적으로 기술해 주세요. (965자)")
label_body("", "**강점 1 — 소재를 아는 부품 영업맨.** OE 지원자 대부분은 기계·전자·산업공학 전공자로 부품의 ‘구조’를 압니다. 저는 부품을 구성하는 ‘소재’의 글로벌 수급과 가격 메커니즘을 압니다. OE 영업의 핵심인 원가분석 기반 견적에서, 부품 원가의 큰 몫을 차지하는 소재(고무·폴리우레탄·강재)의 국제 시황을 실시간으로 읽는다는 건 다른 수준의 협상력입니다. 실제로 PPG 원료 PO의 시황이 폴리우레탄 가격을 거쳐 현가 부품 원가에 어떤 시차로 전이되는지 데이터로 설명할 수 있습니다. “지금이 선구매 타이밍”이라는 논리를 객관적 근거로 제시하는 영업, 그것이 저입니다. 완성차 구매팀과의 원가 협상 테이블에서 이 시황 감각은 비대칭적 무기가 됩니다.")
label_body("", "**강점 2 — 3년간 B2B 풀사이클을 혼자 돌린 신입.** 대다수 신입의 경험은 인턴십·학부 프로젝트입니다. 저는 3년간 영업 전 과정을 독립 수행했습니다. 인도 개척이 대표 사례입니다. 통관 데이터로 시장 규모를 파악하고, 잠재 고객 리스트를 추출해 컨택하고, 샘플 테스트부터 첫 PO까지 6개월의 구매 프로세스를 관리했습니다. 첫 선적 후 품질 이슈가 터지자 기술팀과 협업해 2주 만에 원인을 규명·대응하고 계약을 지켰습니다. 이 사이클은 OE의 RFQ→수주→개발→양산과 구조적으로 같아, 즉시 전력 투입이 가능합니다.")
label_body("", "**약점 — 조향/현가 공학 지식 부족.** 솔직히 R-MDPS 구조나 현가 기구학 설계에 대한 이해는 초보 수준입니다. 기계공학 전공자 대비 기술적 깊이가 부족합니다. 그러나 같은 약점을 이미 극복한 경험이 있습니다. PPG 첫해 화학 비전공자로 3개월을 독학해 1년 만에 팀 매출 20% 담당자가 됐습니다. 지금도 자동차공학·현가설계 강의를 듣고 있고, 입사 후 현장 기술교육을 자청하겠습니다. 오히려 공학 배경이 없기에 엔지니어의 언어가 아닌 ‘고객의 언어’(비즈니스 임팩트)로 구매·원가 담당자와 소통할 수 있다는 점은 협상 테이블에서 강점이 됩니다.")

# ============ 경력기술서 ============
doc.add_page_break()
title("경력기술서")

heading("1. 경력 개요", before=4)
kv("회사", "금호석유화학(주)")
kv("부서", "PPG사업부 해외영업팀")
kv("기간", "2023.01 ~ 현재 (3년차)")
kv("직위", "사원")
kv("직무", "글로벌 B2B 화학소재 해외영업")

heading("2. 주요 담당 업무")
body("아래 업무는 OE 영업의 RFQ·원가분석·프로젝트 관리·이슈 해결과 구조적으로 유사합니다.", size=10, after=8)
for k, v in [
    ("해외영업 총괄", "PPG 제품 일본·인도·유럽·남미 시장 담당"),
    ("시장조사", "통관 데이터 분석 기반 신규 시장 발굴 및 경쟁사 모니터링"),
    ("견적·가격협상", "PO 시황·환율·물류비 연동 원가 분석 후 고객사 견적 제출"),
    ("계약 관리", "Incoterms, L/C, CI·PL·BL 등 수출입 서류 및 외국환 거래"),
    ("고객 관리", "4개 문화권 거래처 동시 대응 (클레임·납기·품질 이슈)"),
    ("프로세스 혁신", "AI·VBA 감사 자동화(1개월→3일), 규제대응 50% 개선"),
]:
    kv(k, v)

heading("3. 핵심 프로젝트 (STAR)")
projects = [
    ("일본 물량 2배 확대", "경쟁사 철수", "점유율 확보",
     "경쟁사 고객리스트 확보 → 품질·납기 제안서 발송 → 첫 달 3곳 거절 → 사유 분석 후 ‘공급 안정성’ 중심으로 제안 피벗",
     "6개월 만에 42톤 → 84톤 (2배)"),
    ("위기 속 가격 인상", "이란-미국 갈등으로 중동발 PO 공급 불안", "장기 계약 갱신",
     "글로벌 PO 플랜트 가동률 데이터로 ‘3개월 내 공급 부족’ 보고서 작성 → 고객사 선제 공유 → 장기계약 유리 논리로 설득",
     "톤당 +$600 인상 조건 갱신"),
    ("인도 신규 개척", "시장 정보 부재", "진입 전략 수립",
     "잠재 고객 리스트 추출·컨택 → 샘플 테스트부터 첫 PO까지 6개월 구매 프로세스 관리 → 첫 선적 품질 이슈 시 기술팀 협업 2주 내 대응",
     "신규 매출처 확보, 계약 유지"),
    ("감사 자료 자동화", "감사 자료 준비에 1개월 소요", "업무 효율화",
     "AI·VBA로 1,000건+ 데이터 자동 분류·정리 프로세스 구축",
     "준비 기간 1개월 → 3일 (90% 단축)"),
]
for name, s, t, a, r in projects:
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(8); pf.space_after = Pt(3); pf.line_spacing = 1.4
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    kfont(p.add_run(f"[{name}]"), 11, True, NAVY)
    for lab, val in [("상황(S)", s), ("과제(T)", t), ("행동(A)", a), ("결과(R)", r)]:
        kv(lab, val)

heading("4. OE 영업 역량 매핑")
for have, oe in [
    ("PO 시황 기반 원가분석·견적", "RFQ 원자재 가격 예측 기반 견적"),
    ("4개 문화권 고객 동시 관리", "글로벌 OE 고객 네트워크 관리"),
    ("가격 인상·인하 방어 협상", "완성차 구매팀 가격 협상 대응"),
    ("감사 자동화·규제대응 개선", "수주 후 사내 유관부서 리드"),
    ("일본어 + 현지 비즈니스 관행", "일본계 OE 고객 영업 강점"),
]:
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.left_indent = Pt(14); pf.space_after = Pt(3); pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    kfont(p.add_run("• "), 11)
    kfont(p.add_run(have), 11, True)
    kfont(p.add_run(f"  →  {oe}"), 11)

out = "/Users/sam/Desktop/brain/자소서/홍석진 취업/2026_자소서결과물/현대모비스_제출용.docx"
doc.save(out)
print("saved:", out)
