# -*- coding: utf-8 -*-
import re

# ===================== 자소서 본문 (정직 버전, 작위 제거) =====================

Q1_TITLE = "[문항 1] 지원동기와 입사 후 회사에서 이루고 싶은 꿈은 무엇인가요?"
Q1 = """제가 영업의 본질을 가장 강하게 체감한 것은 위기를 기회로 바꾼 협상에서였습니다. 이란과 미국의 갈등으로 중동산 원료(PO) 공급에 불안이 커지던 때, 저는 이 시황 리스크를 단가 협상의 근거로 활용했습니다. 공급 불확실성이 길어질수록 안정적 물량 확보가 고객에게 더 중요해진다는 논리로 거래처를 설득해, 톤당 600달러 인상 조건으로 계약을 갱신했습니다. 단순히 값을 올린 것이 아니라, 시장 상황을 함께 읽어 양측 모두를 납득시킨 협상이었습니다. 이때 깨달았습니다. 원료의 원가와 국제 시황을 읽어내는 힘이 곧 협상력이라는 것을요.

이 경험이 제 지원 동기의 출발점입니다. 부품 원가의 큰 몫은 결국 고무·강재·수지 같은 소재이고, 저는 3년간 금호석유화학에서 그 소재의 시황과 원가 구조를 다뤄왔습니다. 소재 영업에서 쌓은 원가 기반 협상력을 자동차 부품 OE 영업으로 확장한다면, 단순히 가격을 제시하는 영업을 넘어 원가의 근거로 설득하는 영업을 할 수 있다고 확신합니다. 소재를 아는 사람이 부품을 판다는 것, 그것이 제가 그리는 차별점입니다.

현대모비스 조향·현가 OE 영업을 택한 이유는 직무 구조 때문입니다. 이 직무는 단순 세일즈가 아니라 RFQ 대응·원가분석·견적·가격협상부터 수주 후 개발 총괄, 양산 안정화까지 프로젝트 전 사이클을 주도하는 자리입니다. 시장조사부터 견적·협상·계약·납품까지 직접 돌려온 제 B2B 영업 경험과 구조가 같습니다. 게다가 조향은 전동식(R-MDPS), 현가는 전자제어 서스펜션으로 전장화가 진행 중이라 성장 시장이라는 점도 매력이었습니다.

입사 후 목표는 단계적입니다. 단기에는 현대차·기아 해외공장 담당으로 OE 영업 사이클을 체화하고, 영어 협상력과 데이터 분석력을 살려 아시아·북미 신규 고객 발굴에 도전하겠습니다. 중기에는 전동화 플랫폼을 맡는 글로벌 OE 영업 담당으로 성장하겠습니다. 장기적으로는 소재 원가까지 읽어내는 영업이라는 강점으로 현대모비스의 해외 OE 매출 확대를 이끄는 사람이 되겠습니다."""

Q2_TITLE = "[문항 2] 지원한 직무를 위해 필요한 역량은 무엇이라고 생각하며, 이 역량을 갖추기 위한 노력이나 자신만의 특별한 경험을 작성해 주세요."
Q2 = """이 직무에 필요한 역량은 논리적 분석력, 진취적 추진력, 조율·협상력, 그리고 글로벌 커뮤니케이션이라고 생각합니다. 저는 이 네 가지를 3년간의 해외영업에서 숫자로 증명해 왔습니다.

**논리적 분석력.** 인도는 회사에 거래선이 전무한 시장이었습니다. 인도가 중국산 제품에 반덤핑 관세를 부과한다는 정보를 기회로 보고, 감이 아닌 데이터로 접근했습니다. 인도·미국 통관 데이터 10만 건을 R과 파이썬으로 전처리해, HS코드가 제각각인 자료에서 누가 경쟁사 제품을 사는지 역추적했습니다. Tokenization으로 송하인·수하인 데이터를 정제한 끝에 현지 유통업체 5곳을 추려내, 신규 진입의 발판을 만들었습니다. 막연한 추정이 아니라 실제 수입 기록에 근거한 분석이었습니다.

**진취적 추진력.** 일본 거래처의 물량이 정체돼 있을 때, 주문을 기다리는 대신 고객의 생산 계획부터 파고들었습니다. 증설 시점과 수요 변화를 읽어 선제적으로 추가 공급을 제안했고, 경쟁사 대비 납기·품질 안정성을 근거로 거래 비중을 키웠습니다. 그 결과 담당 물량을 42톤에서 84톤으로, 1년이 안 돼 두 배로 늘렸습니다. 기다리지 않고 먼저 움직이는 것이 제 영업입니다.

**조율·협상력.** 일본 자동차 내장재 고객사와의 거래에서 chemSHERPA·ELV 등 환경규제 서류 대응이 반복 지연돼, 한 고객사는 물량을 줄이며 거래 중단까지 언급했습니다. 원인은 메일에만 기댄 협업이었습니다. 품질보증팀과 영업팀을 잇는 실시간 소통 채널과 규제 대응 DB를 직접 만들어, 대응 시간을 최대 50% 줄이고 일본 매출을 전년 대비 끌어올렸습니다.

**글로벌 커뮤니케이션.** 저는 일본·인도·유럽·남미 4개 문화권 거래처를 영어(OPIc AL)로 동시에 관리해 왔습니다. 상관습도 의사결정 속도도 다른 상대를 오해 없이 끌고 가려면 언변보다 근거가 중요했습니다. 모든 제안을 데이터로 정리해 전달한 이 방식이, 글로벌 OE 고객을 상대하는 현대모비스의 무대에서 그대로 통할 것이라 믿습니다."""

Q3_TITLE = "[문항 3] 다른 지원자 대비 본인만의 ‘차별화된 강점’과 ‘보완해야 할 약점’에 대해 사례를 들어 구체적으로 기술해 주세요."
Q3 = """**강점 1 — 숫자로 버티는 협상.** 저는 협상 테이블에서 감으로 밀어붙이지 않습니다. 근거를 먼저 챙기는 편입니다. 브라질 거래처가 톤당 30달러를 깎아달라고 했을 때도, 바로 응수하기보다 시장부터 다시 들여다봤습니다. 마침 원료 가격은 오르는 중이었고, 미국 쪽 경쟁사 공장 한 곳은 화재로 가동을 멈춘 상태였습니다. 가격을 내릴 때가 아니라 오히려 공급이 줄고 있다는 사실을 들이밀어야 하는 국면이었던 겁니다. 이 근거로 거래처를 설득해 인하 없이 기존 조건을 지켰습니다. 반대로 공급이 불안할 때는 같은 논리를 거꾸로 써서 단가를 올린 적도 있습니다. 결국 제가 파는 건 가격이 아니라 ‘지금 이 가격이 맞다’는 납득이라고 생각합니다.

**강점 2 — 일하는 방식을 바꾸는 사람.** 외국환거래 감사가 떨어졌을 때, 자금팀이 5년 치 수출 자료를 요청했습니다. 1,000건이 넘는 거래의 PL과 BL을 손으로 하나하나 찾으면 한 달은 잡아야 하는 일이었습니다. 며칠 붙잡아 보니 결국 같은 패턴의 반복이더군요. 그래서 자료를 끌어안는 대신 AI와 VBA를 익혀 검색과 분류 자체를 자동화했습니다. 한 달짜리 일이 3일로 끝났습니다. 오래 붙들고 있는 끈기보다, 일하는 방식 자체를 바꾸는 쪽이 제 성향에 가깝습니다.

**약점 — 조향·현가 공학 지식.** 솔직히 조향과 현가의 공학적 원리는 아직 깊지 못합니다. 전공이 언어와 국제학이라, 처음 화학 제품을 맡았을 때도 제품 구조나 원가 흐름이 영 낯설었습니다. 그때 택한 방법은 단순했습니다. 고객을 설득하려면 말솜씨보다 제품을 먼저 알아야 한다고 봤고, PO 가격과 PPG 용도, 원료 시황과 판가 구조를 업무 속에서 끈질기게 파고들었습니다. 덕분에 시황과 원가를 근거로 가격을 지키고 조건을 끌어올릴 수 있었습니다. 조향·현가도 다르지 않다고 봅니다. 입사 후 제품 구조와 고객별 RFQ 기준, 원가 항목, 개발 프로세스부터 빠르게 익혀, 엔지니어와 구매 담당자 양쪽의 언어를 함께 쓰는 OE 영업 담당자가 되겠습니다."""

def count(s):
    return len(s.replace("**", ""))

counts = {}
for tag, q in [("문항1", Q1), ("문항2", Q2), ("문항3", Q3)]:
    c = count(q)
    counts[tag] = c
    ok = "OK" if 950 <= c <= 990 else ("LOW" if c < 950 else "HIGH")
    print(f"{tag}: {c}자  [{ok}]")

BASE = "/Users/sam/Desktop/brain/자소서/홍석진 취업/2026_자소서결과물/"

# ===================== 경력기술서 데이터 (표 없이 텍스트) =====================
OVERVIEW = [
    ("회사", "금호석유화학(주)"),
    ("부서", "PPG사업부 해외영업팀"),
    ("기간", "2023.01 ~ 현재 (3년차)"),
    ("직위", "사원"),
    ("직무", "글로벌 B2B 화학소재 해외영업"),
]
DUTIES = [
    ("해외영업 총괄", "PPG 제품 일본·인도·유럽·남미 시장 담당"),
    ("시장조사", "데이터 분석 기반 신규 시장 발굴 및 경쟁사 모니터링"),
    ("견적·가격협상", "PO 시황·환율·물류비 연동 원가 분석 후 고객사 견적 제출"),
    ("계약 관리", "Incoterms, L/C, CI·PL·BL 등 수출입 서류 및 외국환 거래"),
    ("고객 관리", "4개 문화권 거래처 동시 대응 (클레임·납기·품질 이슈)"),
    ("프로세스 혁신", "AI·VBA 감사 자동화(1개월→3일), 규제대응 50% 개선"),
]
PROJECTS = [
    ("인도 신규 시장 개척", "회사에 인도 거래선 전무, 인도의 중국산 반덤핑 관세 부과", "신규 거래선 발굴",
     "인도·미국 통관 데이터 10만 건을 R·파이썬으로 전처리 → Tokenization으로 송하인·수하인 정제 → 유통업체 5곳 도출 → ‘Dow와 경쟁 안 된다’는 거절에도 끈질긴 컨택",
     "현지 거래선 최초 발굴, 중국 빈자리 진입 기회 확보"),
    ("일본 규제대응 프로세스 개선", "EU 환경규제(chemSHERPA·ELV) 서류 지연 → 고객 물량 축소·거래 중단 언급", "신뢰 회복·매출 방어",
     "지연 근본원인(메일 위주 협업) 분석 → 규제 학습·대응 DB 구축 → 품질보증팀·영업팀 실시간 소통 채널 개설",
     "대응 시간 최대 50% 단축, 일본 매출 전년 대비 증가"),
    ("일본 물량 2배 확대", "일본 거래처 물량 정체", "담당 물량 확대",
     "고객 생산 계획·증설 시점 분석 → 선제적 추가 공급 제안 → 경쟁사 대비 납기·품질 안정성 근거로 거래 비중 확대",
     "담당 물량 42톤 → 84톤 (1년 내 2배)"),
    ("중동 단가 인상 협상", "이란-미국 갈등으로 중동산 원료(PO) 공급 불안", "계약 단가 방어·인상",
     "시황 리스크를 근거로 ‘안정 물량 확보가 고객에 유리’ 논리 구성 → 거래처 설득",
     "톤당 +$600 인상 조건 갱신 (브라질 -$30 인하 요구는 방어)"),
    ("감사 자료 자동화", "감사 자료 준비에 1개월 소요", "업무 효율화",
     "1,000건+ 서류를 AI·VBA로 자동 분류·정리하는 프로세스 구축",
     "준비 기간 1개월 → 3일 (90% 단축)"),
]
MAPPING = [
    ("PO 시황 기반 원가분석·견적", "RFQ 원자재 가격 예측 기반 견적"),
    ("4개 문화권 고객 동시 관리", "글로벌 OE 고객 네트워크 관리"),
    ("가격 인상·인하 방어 협상", "완성차 구매팀 가격 협상 대응"),
    ("감사 자동화·규제대응 개선", "수주 후 사내 유관부서 리드"),
    ("통관 데이터 10만 건 분석 경험", "경쟁사 입찰 분석·신규 고객 발굴"),
]

# ===================== MD 출력 =====================
md = []
md.append("# 현대모비스 Global OE 조향/현가 영업 — 자기소개서 및 경력기술서\n")
for title_, q, tag in [(Q1_TITLE, Q1, "문항1"), (Q2_TITLE, Q2, "문항2"), (Q3_TITLE, Q3, "문항3")]:
    md.append(f"## {title_} *({counts[tag]}자)*\n")
    md.append(q + "\n")
    md.append("---\n")
md.append("# 경력기술서\n")
md.append("## 1. 경력 개요\n")
for k, v in OVERVIEW:
    md.append(f"- **{k}** : {v}")
md.append("\n## 2. 주요 담당 업무\n")
for k, v in DUTIES:
    md.append(f"- **{k}** : {v}")
md.append("\n## 3. 핵심 프로젝트 (STAR)\n")
for name, s, t, a, r in PROJECTS:
    md.append(f"**[{name}]**")
    md.append(f"- 상황(S) : {s}")
    md.append(f"- 과제(T) : {t}")
    md.append(f"- 행동(A) : {a}")
    md.append(f"- 결과(R) : {r}\n")
md.append("## 4. OE 영업 역량 매핑\n")
for h, o in MAPPING:
    md.append(f"- **{h}**  →  {o}")
md.append("\n*작성일자: 2026-06-14*")
with open(BASE + "현대모비스_제출용.md", "w", encoding="utf-8") as f:
    f.write("\n".join(md))
print("md saved")

# ===================== DOCX 출력 =====================
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING

FONT = "맑은 고딕"
NAVY = RGBColor(0x00, 0x33, 0x66)
doc = Document()
sty = doc.styles["Normal"]; sty.font.name = FONT; sty.font.size = Pt(11)
sty._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

def kf(run, size=11, bold=False, color=None):
    run.font.name = FONT; run.font.size = Pt(size); run.font.bold = bold
    rpr = run._element.get_or_add_rPr(); rfe = rpr.find(qn("w:rFonts"))
    if rfe is None:
        rfe = OxmlElement("w:rFonts"); rpr.append(rfe)
    rfe.set(qn("w:eastAsia"), FONT); rfe.set(qn("w:ascii"), FONT); rfe.set(qn("w:hAnsi"), FONT)
    if color: run.font.color.rgb = color

def ttl(text, size=18):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(4); pf.space_after = Pt(14); pf.line_spacing = 1.2
    kf(p.add_run(text), size, True, NAVY)

def hd(text, size=13, before=14):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(8)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.3
    kf(p.add_run(text), size, True, NAVY)

def inline(p, text, size=11):
    for seg in re.split(r"(\*\*.*?\*\*)", text):
        if seg.startswith("**") and seg.endswith("**"):
            kf(p.add_run(seg[2:-2]), size, True)
        elif seg:
            kf(p.add_run(seg), size)

def para_block(text, size=11):
    # split on blank lines into paragraphs
    for block in text.split("\n\n"):
        p = doc.add_paragraph(); pf = p.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.7
        inline(p, block.strip(), size)

def kv(k, v):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.left_indent = Pt(14); pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.5
    kf(p.add_run("• "), 11); kf(p.add_run(f"{k} : "), 11, True); kf(p.add_run(v), 11)

def divider():
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    kf(p.add_run("─" * 28), 9, color=RGBColor(0xAA, 0xAA, 0xAA))

ttl("현대모비스 Global OE 조향/현가 영업 — 자기소개서 및 경력기술서")
for title_, q, tag, first in [(Q1_TITLE, Q1, "문항1", True), (Q2_TITLE, Q2, "문항2", False), (Q3_TITLE, Q3, "문항3", False)]:
    hd(f"{title_} ({counts[tag]}자)", before=8 if first else 14)
    para_block(q)
    divider()

doc.add_page_break()
ttl("경력기술서")
hd("1. 경력 개요", before=4)
for k, v in OVERVIEW: kv(k, v)
hd("2. 주요 담당 업무")
para_block("아래 업무는 OE 영업의 RFQ·원가분석·프로젝트 관리·이슈 해결과 구조적으로 유사합니다.")
for k, v in DUTIES: kv(k, v)
hd("3. 핵심 프로젝트 (STAR)")
for name, s, t, a, r in PROJECTS:
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(8); pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.4
    kf(p.add_run(f"[{name}]"), 11, True, NAVY)
    for lab, val in [("상황(S)", s), ("과제(T)", t), ("행동(A)", a), ("결과(R)", r)]:
        kv(lab, val)
hd("4. OE 영업 역량 매핑")
for h, o in MAPPING:
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.left_indent = Pt(14); pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.5
    kf(p.add_run("• "), 11); kf(p.add_run(h), 11, True); kf(p.add_run(f"  →  {o}"), 11)

doc.save(BASE + "현대모비스_제출용.docx")
print("docx saved")
