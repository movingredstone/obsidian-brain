# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING

FONT = "맑은 고딕"
NAVY = RGBColor(0x00, 0x33, 0x66)

doc = Document()
st = doc.styles["Normal"]
st.font.name = FONT
st.font.size = Pt(11)
st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

def set_kfont(run, size=11, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        from docx.oxml import OxmlElement
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:eastAsia"), FONT); rf.set(qn("w:ascii"), FONT); rf.set(qn("w:hAnsi"), FONT)
    if color: run.font.color.rgb = color

def heading(text, size=14):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12); pf.space_after = Pt(8)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.3
    r = p.add_run(text); set_kfont(r, size, True, NAVY)

def bullet(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Pt(14); pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.5
    r0 = p.add_run("• "); set_kfont(r0, 11)
    r = p.add_run(text); set_kfont(r, 11)

# Title
p = doc.add_paragraph(); pf = p.paragraph_format; pf.space_after = Pt(14); pf.line_spacing = 1.2
r = p.add_run("경력기술서"); set_kfont(r, 18, True, NAVY)

# 경력 개요
heading("1. 경력 개요")
bullet("금호석유화학(주) PPG사업부 해외영업팀 / 사원 / 2023.01 ~ 현재 (3년차)")
bullet("글로벌 B2B 화학소재(PPG) 해외영업 담당 — 일본·인도·유럽·남미 시장")

# 한 일 나열
heading("2. 담당 업무")
duties = [
    "PPG 제품 해외영업 총괄 — 일본·인도·유럽·남미 4개 시장 담당",
    "시황·환율·물류비를 반영한 원가 분석 후 고객사 견적 제출 및 가격 협상",
    "Incoterms, L/C, CI·PL·BL 등 수출입 서류 작성 및 외국환 거래 처리",
    "4개 문화권 거래처 동시 관리 — 클레임·납기·품질 이슈 대응",
    "신규 시장·고객 발굴 및 경쟁사 동향 모니터링",
    "AI·VBA 활용 감사 자료 준비 자동화로 1개월 → 3일 단축",
    "화학 비전공 입사 후 독학으로 1년 만에 팀 매출 20% 책임 담당자로 성장",
]
for d in duties:
    bullet(d)

# 주요 성과 나열
heading("3. 주요 성과")
results = [
    "일본 시장 물량 6개월 만에 42톤 → 84톤으로 2배 확대",
    "중동 공급 불안기에 장기계약 갱신 — 톤당 +$600 인상",
    "브라질 거래처의 톤당 -$30 인하 요구를 시황 데이터로 방어",
    "인도 신규 시장 진입 및 신규 매출처 확보",
    "감사 자료 준비 프로세스 자동화로 업무 시간 90% 단축",
]
for r in results:
    bullet(r)

out = "/Users/sam/Desktop/brain/자소서/홍석진 취업/2026_자소서결과물/현대모비스_경력기술서_복붙용.docx"
doc.save(out)
print("saved:", out)
